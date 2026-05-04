"""
Generador de informes Word (.docx) con gráficos Plotly embebidos.

Convierte las figuras interactivas de Plotly a imágenes PNG y las inserta
en un documento Word editable con formato institucional de la Policía de Tucumán.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Optional

import plotly.graph_objects as go

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import kaleido  # noqa: F401 – solo verificar que esté disponible
    HAS_KALEIDO = True
except ImportError:
    HAS_KALEIDO = False


def _check_deps():
    """Lanza un error claro si faltan dependencias."""
    missing = []
    if not HAS_DOCX:
        missing.append("python-docx")
    if not HAS_KALEIDO:
        missing.append("kaleido")
    if missing:
        raise ImportError(
            f"Para exportar informes se necesitan: {', '.join(missing)}. "
            f"Instálalas con: pip install {' '.join(missing)}"
        )


def _fig_to_png_bytes(fig: go.Figure, width: int = 1100, height: int = 600) -> bytes:
    """Convierte una figura Plotly a bytes PNG usando kaleido."""
    return fig.to_image(format="png", width=width, height=height, scale=2)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Agrega un encabezado con formato institucional."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x11, 0x21, 0x3C)


def _add_metric_table(doc: Document, metricas: dict[str, str]):
    """Crea una tabla horizontal con métricas tipo KPI."""
    if not metricas:
        return

    table = doc.add_table(rows=2, cols=len(metricas))
    table.style = "Light Grid Accent 1"

    for idx, (label, valor) in enumerate(metricas.items()):
        # Header row
        cell_header = table.rows[0].cells[idx]
        cell_header.text = label
        for paragraph in cell_header.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.bold = True

        # Value row
        cell_value = table.rows[1].cells[idx]
        cell_value.text = str(valor)
        for paragraph in cell_value.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    doc.add_paragraph()  # Espaciado


class WordReportBuilder:
    """
    Construye un documento Word con gráficos y métricas del dashboard.

    Uso típico:
        builder = WordReportBuilder(titulo="Informe Delictual")
        builder.add_section("Análisis General")
        builder.add_metricas({"Total Delitos": "1,234", "Jurisdicciones": "45"})
        builder.add_chart(fig_plotly, "Distribución por Modalidad")
        buffer = builder.build()
        # → buffer es un BytesIO listo para st.download_button
    """

    def __init__(
        self,
        titulo: str = "INFORME DELICTUAL",
        subtitulo: str = "Policía de la Provincia de Tucumán",
        orientacion: str = "vertical",
    ):
        _check_deps()
        self.doc = Document()
        self._setup_page(orientacion)
        self._add_header(titulo, subtitulo)

    def _setup_page(self, orientacion: str):
        """Configura márgenes y orientación del documento."""
        section = self.doc.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        if orientacion == "horizontal":
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width = section.page_height
            new_height = section.page_width
            section.page_width = new_width
            section.page_height = new_height

    def _add_header(self, titulo: str, subtitulo: str):
        """Agrega el encabezado institucional al documento."""
        # Título principal
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(titulo)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x11, 0x21, 0x3C)

        # Subtítulo institucional
        sub_para = self.doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitulo)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Fecha de generación
        fecha_para = self.doc.add_paragraph()
        fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_run = fecha_para.add_run(
            f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M hs')}"
        )
        fecha_run.font.size = Pt(10)
        fecha_run.font.italic = True
        fecha_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Línea separadora
        self.doc.add_paragraph("─" * 80)

    def add_section(self, titulo: str, descripcion: str = ""):
        """Agrega una sección con título y descripción opcional."""
        _add_heading(self.doc, titulo, level=2)
        if descripcion:
            para = self.doc.add_paragraph(descripcion)
            para.style = self.doc.styles["Normal"]
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def add_metricas(self, metricas: dict[str, str]):
        """Agrega una tabla de métricas al documento."""
        _add_metric_table(self.doc, metricas)

    def add_chart(
        self,
        fig: go.Figure,
        titulo: str = "",
        descripcion: str = "",
        width: int = 1100,
        height: int = 600,
    ):
        """
        Convierte una figura Plotly a imagen y la inserta en el documento.

        Parameters
        ----------
        fig : go.Figure
            La figura Plotly a convertir.
        titulo : str
            Título que aparecerá arriba del gráfico en el Word.
        descripcion : str
            Texto descriptivo debajo del gráfico.
        width, height : int
            Dimensiones en píxeles para la imagen exportada.
        """
        if titulo:
            _add_heading(self.doc, titulo, level=3)

        # Convertir la figura a PNG
        try:
            # Crear una copia para ajustar colores a "impresión"
            fig_export = go.Figure(fig)
            fig_export.update_layout(
                paper_bgcolor="white",
                plot_bgcolor="#f8f9fa",
                font=dict(color="#111111", size=14),
                title_font=dict(color="#11213c", size=18),
            )
            fig_export.update_xaxes(
                tickfont=dict(color="#333333"),
                title_font=dict(color="#333333"),
                gridcolor="#dddddd",
            )
            fig_export.update_yaxes(
                tickfont=dict(color="#333333"),
                title_font=dict(color="#333333"),
                gridcolor="#dddddd",
            )

            img_bytes = _fig_to_png_bytes(fig_export, width=width, height=height)

            # Insertar la imagen en el documento
            img_stream = io.BytesIO(img_bytes)
            # Calcular ancho disponible (A4 con márgenes)
            available_width = Inches(6.5)
            self.doc.add_picture(img_stream, width=available_width)

            # Centrar la imagen
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            error_para = self.doc.add_paragraph(
                f"[Error al exportar gráfico: {e}]"
            )
            for run in error_para.runs:
                run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

        if descripcion:
            desc_para = self.doc.add_paragraph(descripcion)
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in desc_para.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        self.doc.add_paragraph()  # Espaciado

    def add_text(self, texto: str, bold: bool = False):
        """Agrega un párrafo de texto libre al documento."""
        para = self.doc.add_paragraph()
        run = para.add_run(texto)
        run.font.size = Pt(11)
        run.font.bold = bold

    def add_page_break(self):
        """Inserta un salto de página."""
        self.doc.add_page_break()

    def build(self) -> io.BytesIO:
        """
        Genera el documento Word y lo retorna como BytesIO.

        Returns
        -------
        io.BytesIO
            Buffer con el archivo .docx listo para descarga.
        """
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
